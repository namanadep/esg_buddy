"""Convert PROJECT_TECHNICAL_OVERVIEW.md to a formatted .docx file."""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

MD_PATH = Path(__file__).parent / "PROJECT_TECHNICAL_OVERVIEW.md"
DOCX_PATH = Path(__file__).parent / "PROJECT_TECHNICAL_OVERVIEW.docx"


def set_cell_shading(cell, color_hex):
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): color_hex,
    })
    shading.append(shd)


def add_formatted_paragraph(doc, text, style="Normal", bold=False, italic=False,
                            alignment=None, font_size=None, space_after=None,
                            font_color=None):
    p = doc.add_paragraph(style=style)
    if alignment is not None:
        p.alignment = alignment

    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            run = p.add_run(part)
            run.bold = bold

        run.italic = italic
        if font_size:
            run.font.size = Pt(font_size)
        if font_color:
            run.font.color.rgb = font_color

    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)

    return p


def add_table_from_lines(doc, header_line, data_lines):
    headers = [c.strip() for c in header_line.strip("|").split("|")]
    rows = []
    for dl in data_lines:
        cells = [c.strip() for c in dl.strip("|").split("|")]
        rows.append(cells)

    num_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=num_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        set_cell_shading(cell, "D9E2F3")

    for r_idx, row_data in enumerate(rows):
        for c_idx in range(min(num_cols, len(row_data))):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = row_data[c_idx]
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(10)

    doc.add_paragraph("")


def convert():
    md_text = MD_PATH.read_text(encoding="utf-8")
    lines = md_text.split("\n")

    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.15

    for level in range(1, 5):
        hs = doc.styles[f"Heading {level}"]
        hs.font.name = "Calibri"
        hs.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped or stripped == "---":
            i += 1
            continue

        if stripped.startswith("# ") and not stripped.startswith("## "):
            heading_text = stripped[2:].strip()
            p = doc.add_heading(heading_text, level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        if stripped.startswith("## "):
            heading_text = stripped[3:].strip()
            doc.add_heading(heading_text, level=2)
            i += 1
            continue

        if stripped.startswith("### "):
            heading_text = stripped[4:].strip()
            doc.add_heading(heading_text, level=3)
            i += 1
            continue

        if stripped.startswith("| ") and i + 1 < len(lines):
            next_line = lines[i + 1].strip() if i + 1 < len(lines) else ""
            if re.match(r"^\|[\s\-:|]+\|$", next_line):
                header_line = stripped
                data_lines = []
                j = i + 2
                while j < len(lines) and lines[j].strip().startswith("|"):
                    data_lines.append(lines[j].strip())
                    j += 1
                add_table_from_lines(doc, header_line, data_lines)
                i = j
                continue

        if stripped.startswith("- **") or stripped.startswith("- "):
            text = stripped[2:].strip()
            add_formatted_paragraph(doc, text, style="List Bullet")
            i += 1
            continue

        if re.match(r"^\d+\.\s", stripped):
            text = re.sub(r"^\d+\.\s*", "", stripped)
            add_formatted_paragraph(doc, text, style="List Number")
            i += 1
            continue

        add_formatted_paragraph(doc, stripped, space_after=6)
        i += 1

    doc.save(str(DOCX_PATH))
    print(f"Saved: {DOCX_PATH}")


if __name__ == "__main__":
    convert()
